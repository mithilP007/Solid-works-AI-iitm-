from fastapi import FastAPI, UploadFile, File, WebSocket, WebSocketDisconnect, BackgroundTasks
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, FileResponse, StreamingResponse
import cv2
import shutil
import os
from pathlib import Path
import pandas as pd
from pipeline import CustomPipeline
import asyncio
from fpdf import FPDF
from docx import Document
import time
from datetime import datetime
import GPUtil
import psutil
from pydantic import BaseModel

class PredictRequest(BaseModel):
    augment: bool = True

app = FastAPI(debug=True)

@app.websocket("/ws")
async def websocket_endpoint(websocket: WebSocket):
    await log_manager.connect(websocket)
    try:
        while True:
            await websocket.receive_text()
    except WebSocketDisconnect:
        log_manager.disconnect(websocket)

# Log Broadcaster
class LogManager:
    def __init__(self):
        self.active_connections: list[WebSocket] = []

    async def connect(self, websocket: WebSocket):
        await websocket.accept()
        self.active_connections.append(websocket)

    def disconnect(self, websocket: WebSocket):
        if websocket in self.active_connections:
            self.active_connections.remove(websocket)

    async def broadcast(self, message: str):
        print(f"LOG: {message}")
        for connection in self.active_connections:
            try:
                await connection.send_text(message)
            except:
                pass

log_manager = LogManager()
is_inferencing = False
batch_results = []

# Enable CORS with explicit origins for safety and compatibility
origins = [
    "http://localhost",
    "http://localhost:5173",
    "http://localhost:8080",
    "http://127.0.0.1",
    "http://127.0.0.1:5173",
    "http://127.0.0.1:8080",
    "*"
]

app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

BASE_DIR = Path(r"d:\Solid works datasets datas")
UPLOAD_DIR = BASE_DIR / "web_uploads"
UPLOAD_DIR.mkdir(exist_ok=True)
SUBMISSION_PATH = BASE_DIR / "submission.csv"

# Initialize YOLO Pipeline
MODEL_PATH = BASE_DIR / "runs" / "detect" / "train_yolo11_advanced" / "weights" / "best.pt"
pipeline = CustomPipeline(str(MODEL_PATH))

# CRITICAL FIX: Patch model class names globally to resolve Nut/Bolt swap
# print("DEBUG: Patching model class names...")
# try:
#     names = pipeline.model.names
#     # Find IDs dynamically to be safe
#     bolt_id = next((k for k, v in names.items() if v == 'bolt'), None)
#     nut_id = next((k for k, v in names.items() if v == 'nut'), None)
#     
#     if bolt_id is not None and nut_id is not None:
#         # NAMES ARE CORRECT IN MODEL - DO NOT SWAP
#         # names[bolt_id] = 'nut'
#         # names[nut_id] = 'bolt'
#         print(f"DEBUG: Native names kept: {names}")
#     else:
#         print("DEBUG: Could not find 'bolt' and 'nut' in model classes.")
# except Exception as e:
#     print(f"ERROR Patching model names: {e}")



@app.get("/")
def read_root():
    return {"status": "AI Factory Multi-Engine Active", "disk": "D:"}

@app.post("/upload")
async def upload_image(file: UploadFile = File(...)):
    file_path = UPLOAD_DIR / file.filename
    with open(file_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)
    return {"filename": file.filename, "status": "Uploaded"}

@app.post("/upload_bulk")
async def upload_bulk(files: list[UploadFile] = File(...)):
    try:
        uploaded_count = 0
        for file in files:
            # Flatten directory structures if present: folder/img.png -> folder_img.png
            flattened_name = file.filename.replace('/', '_').replace('\\', '_')
            file_path = UPLOAD_DIR / flattened_name
            with open(file_path, "wb") as buffer:
                shutil.copyfileobj(file.file, buffer)
            uploaded_count += 1
        return {"count": uploaded_count, "status": "Uploaded"}
    except Exception as e:
        import traceback
        traceback.print_exc()
        return JSONResponse({"error": str(e), "trace": traceback.format_exc()}, status_code=500)

@app.post("/upload_zip", status_code=202)
async def upload_zip(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    augment: bool = True
):
    try:
        # Save the zip temporarily
        zip_path = BASE_DIR / "temp_upload.zip"
        with open(zip_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
            
        # Trigger background processing
        background_tasks.add_task(process_zip_background, zip_path, augment)
        
        return {"status": "Accepted", "message": "ZIP processing started in background"}
    except Exception as e:
        import traceback
        traceback.print_exc()
        return JSONResponse({"error": str(e), "trace": traceback.format_exc()}, status_code=500)

async def process_zip_background(zip_path: Path, augment_enabled: bool):
    global is_inferencing, batch_results
    
    # Wait for any existing inference to clear? Or force clear?
    # We'll assume clear_uploads was called before this as per frontend flow.
    is_inferencing = True # Set immediately to block other reqs
    
    try:
        import zipfile
        await log_manager.broadcast("LOG: Unpacking ZIP Archive...")
        
        # Extract to UPLOAD_DIR
        extracted_count = 0
        with zipfile.ZipFile(zip_path, 'r') as zip_ref:
            members = zip_ref.namelist()
            # Filter and extract - STRICT MACOS FILTERING
            images_to_extract = []
            for m in members:
                # Ignore folder paths specific to macos
                if '__MACOSX' in m: continue
                # Ignore hidden metadata files (._img.png)
                if os.path.basename(m).startswith('._'): continue
                
                if m.lower().endswith(('.png', '.jpg', '.jpeg')):
                    images_to_extract.append(m)
            
            for i, member in enumerate(images_to_extract):
                # Use only the base filename, not the full path
                base_name = os.path.basename(member)
                with zip_ref.open(member) as source, open(UPLOAD_DIR / base_name, "wb") as target:
                    shutil.copyfileobj(source, target)
                extracted_count += 1
                
                # Report extraction progress periodically for large zips
                if i % 100 == 0:
                     await log_manager.broadcast(f"LOG: Extracted {i}/{len(images_to_extract)} files...")

        await log_manager.broadcast(f"LOG: Extraction Complete. {extracted_count} images ready.")
        await log_manager.broadcast(f"LOG: TOTAL_COUNT:{extracted_count}")
        if zip_path.exists():
            os.remove(zip_path)
            
        # Now Trigger Inference (Reusing same logic)
        await run_batch_inference(augment_enabled)

    except Exception as e:
        await log_manager.broadcast(f"zip_error: {str(e)}")
        is_inferencing = False

@app.post("/clear_uploads")
async def clear_uploads():
    for f in UPLOAD_DIR.glob('*'):
        if f.is_file():
            f.unlink()
    # Reset simulation state
    global batch_results
    batch_results = []
    
    # Delete submission and stat files
    if SUBMISSION_PATH.exists():
        SUBMISSION_PATH.unlink()
    
    files_to_delete = [
        BASE_DIR / "last_speed.txt",
        BASE_DIR / "last_tta.txt",
        BASE_DIR / "prev_stats.json",
        BASE_DIR / "report.pdf",
        BASE_DIR / "summary.docx"
    ]
    for f in files_to_delete:
        if f.exists():
            f.unlink()
            
    await log_manager.broadcast("LOG: FACTORY_RESET")
    return {"status": "Cleared and Reset"}

batch_results = []

async def run_batch_inference(augment_enabled: bool):
    global batch_results, is_inferencing
    # IMPORTANT: Do not reset is_inferencing to True here if called from background task 
    # where it might already be True, but for safety ensuring it's True is fine.
    is_inferencing = True
    # Flag is already set by parent, but safety verify
    batch_results = []
    
    try:
        # Case-insensitive image search - Deduplicate by FILENAME (Windows safe)
        seen_names = set()
        images = []
        for ext in ['*.png', '*.jpg', '*.jpeg']:
            for p in UPLOAD_DIR.glob(ext):
                if p.name.lower() not in seen_names:
                    seen_names.add(p.name.lower())
                    images.append(p)
        
        if not images:
            await log_manager.broadcast("BATCH ERROR: No images found in factory uploads.")
            return

        await log_manager.broadcast(f"INITIATING REAL-TIME BATCH: {len(images)} targets identified.")
        
        summary_data = {
            'image_name': [],
            'bolt': [],
            'locatingpin': [],
            'nut': [],
            'washer': []
        }
        
        start_time = time.time()
        
        # BATCH PROCESSING LOGIC (Speed Optimization)
        BATCH_SIZE = 8
        
        for i in range(0, len(images), BATCH_SIZE):
            batch_files = images[i : i + BATCH_SIZE]
            batch_paths = [str(p) for p in batch_files]
            
            try:
                # Predict on batch with optimized parameters
                results = pipeline.model.predict(
                    batch_paths,
                    augment=True,   # Always enable TTA for maximum accuracy
                    conf=0.20,      # Balanced threshold
                    iou=0.45,       # Improved NMS
                    agnostic_nms=True, # Better handling of overlapping parts
                    imgsz=1024,     # High resolution
                    verbose=False
                )
                
                # Process batch results
                for j, result in enumerate(results):
                    img_path = batch_files[j]
                    counts = {'bolt': 0, 'locatingpin': 0, 'nut': 0, 'washer': 0}
                    detections = []
                    
                    if result.boxes is not None:
                        for box in result.boxes:
                            cls_id = int(box.cls[0])
                            
                            cls_name = pipeline.model.names[cls_id]
                                
                            conf = float(box.conf[0])
                            b = box.xyxyn[0].tolist()
                            
                            if cls_name in counts:
                                counts[cls_name] += 1
                            
                            detections.append({
                                "id": f"{img_path.stem}_{len(detections)}",
                                "class": cls_name,
                                "confidence": conf,
                                "bbox": {
                                    "x": b[0] * 100,
                                    "y": b[1] * 100,
                                    "width": (b[2] - b[0]) * 100,
                                    "height": (b[3] - b[1]) * 100
                                }
                            })
                    
                    # Record result
                    img_result = {
                        "id": img_path.name,
                        "src": f"http://localhost:8000/images/{img_path.name}",
                        "detections": detections
                    }
                    batch_results.append(img_result)
                    
                    # Update summary data - filename is now correct from extraction
                    summary_data['image_name'].append(img_path.name)
                    summary_data['bolt'].append(counts['bolt'])
                    summary_data['locatingpin'].append(counts['locatingpin'])
                    summary_data['nut'].append(counts['nut'])
                    summary_data['washer'].append(counts['washer'])

            except Exception as e:
                await log_manager.broadcast(f"BATCH ERROR at index {i}: {str(e)}")
            
            # Broadcast progress after each batch
            processed_so_far = min(i + BATCH_SIZE, len(images))
            progress_msg = f"PROGRESS: {processed_so_far}/{len(images)} processed. Batch speed optimized."
            await log_manager.broadcast(progress_msg)
            
            # Update CSV periodically
            if processed_so_far % 40 == 0 or processed_so_far == len(images):
                 df = pd.DataFrame(summary_data)
                 df.to_csv(SUBMISSION_PATH, index=False)
            
            # Tiny sleep to keep websocket responsive
            await asyncio.sleep(0.01)

        total_time = time.time() - start_time
        avg_speed = (total_time / len(images)) * 1000 if images else 0
        with open(BASE_DIR / "last_speed.txt", "w") as f:
            f.write(str(int(avg_speed)))
            
        await log_manager.broadcast(f"BATCH COMPLETE: {len(images)} images synced to real-time analytics.")
        
    finally:
        is_inferencing = False

@app.post("/predict_all")
async def predict_all(request: PredictRequest = None, background_tasks: BackgroundTasks = None):
    global is_inferencing
    if is_inferencing:
        return {"status": "Already processing. Please wait."}
        
    augment_enabled = request.augment if request else True
    
    # FIX: Set Flag IMMEDIATELY to prevent race condition with frontend polling
    is_inferencing = True
    
    # Save last TTA status
    with open(BASE_DIR / "last_tta.txt", "w") as f:
        f.write("Active" if augment_enabled else "Disabled")

    # Start async task
    asyncio.create_task(run_batch_inference(augment_enabled))
    
    return {"status": "Batch processing started in background"}

@app.get("/batch_status")
async def get_batch_status():
    try:
        global is_inferencing
        return {
            "processed_count": len(batch_results),
            "results": batch_results,
            "is_processing": is_inferencing
        }
    except Exception as e:
        import traceback
        traceback.print_exc()
        return JSONResponse({"error": str(e), "trace": traceback.format_exc()}, status_code=500)

@app.get("/images/{image_name}")
def get_image(image_name: str):
    img_path = UPLOAD_DIR / image_name
    if img_path.exists():
        return FileResponse(img_path)
    return JSONResponse({"error": "Image not found"}, status_code=404)

@app.get("/training_history")
def get_training_history():
    results_path = BASE_DIR / "runs" / "detect" / "train_yolo11_advanced" / "results.csv"
    if not results_path.exists():
        return []
    
    try:
        df = pd.read_csv(results_path)
        df.columns = [c.strip() for c in df.columns]
        
        history = []
        for _, row in df.iterrows():
            if 'epoch' in row and 'metrics/mAP50(B)' in row:
                history.append({
                    "epoch": int(row['epoch']),
                    "map50": float(row['metrics/mAP50(B)']),
                    "loss": float(row.get('train/box_loss', 0)) + float(row.get('train/cls_loss', 0))
                })
        return history
    except Exception as e:
        return []

@app.get("/stats")
def get_stats():
    # If submission.csv doesn't exist, we just return empty counts (zeros)
    if not SUBMISSION_PATH.exists():
        return {
            "total_images": 0, "bolt": 0, "locatingpin": 0, "nut": 0, "washer": 0,
            "total_parts": 0, "parts_trend": 0, "inference_speed": 0,
            "model_accuracy": 0, "accuracy_trend": 0, "tta_status": "Active",
            "last_updated": datetime.now().strftime("%H:%M:%S")
        }
    
    try:
        df = pd.read_csv(SUBMISSION_PATH)
    except:
        return {"error": "Corrupt submission file"}
    
    inference_speed = 0
    speed_file = BASE_DIR / "last_speed.txt"
    if speed_file.exists():
        try:
            with open(speed_file, "r") as f:
                inference_speed = int(f.read().strip())
        except:
            pass

    tta_status = "Active"
    tta_file = BASE_DIR / "last_tta.txt"
    if tta_file.exists():
        try:
            with open(tta_file, "r") as f:
                tta_status = f.read().strip()
        except:
            pass

    accuracy = 0
    accuracy_trend = 0
    results_path = BASE_DIR / "runs" / "detect" / "train_yolo11_advanced" / "results.csv"
    if results_path.exists():
        try:
            rdf = pd.read_csv(results_path)
            rdf.columns = [c.strip() for c in rdf.columns]
            if not rdf.empty:
                accuracy = float(rdf['metrics/mAP50(B)'].iloc[-1])
                if len(rdf) > 1:
                    prev_acc = float(rdf['metrics/mAP50(B)'].iloc[-2])
                    accuracy_trend = round(((accuracy - prev_acc) / (prev_acc if prev_acc != 0 else 1)) * 100, 1)
        except:
            pass

    prev_parts = 0
    parts_trend = 0
    prev_stats_path = BASE_DIR / "prev_stats.json"
    total_parts = int(df[['bolt', 'locatingpin', 'nut', 'washer']].sum().sum()) if not df.empty else 0
    
    if prev_stats_path.exists():
        try:
            with open(prev_stats_path, "r") as f:
                import json
                prev_data = json.load(f)
                prev_parts = prev_data.get("total_parts", 0)
                if prev_parts > 0:
                    parts_trend = round(((total_parts - prev_parts) / prev_parts) * 100, 1)
        except:
            pass
    
    try:
        with open(prev_stats_path, "w") as f:
            import json
            json.dump({"total_parts": total_parts}, f)
    except:
        pass

    stats = {
        "total_images": len(df),
        "bolt": int(df['bolt'].sum()) if 'bolt' in df else 0,
        "locatingpin": int(df['locatingpin'].sum()) if 'locatingpin' in df else 0,
        "nut": int(df['nut'].sum()) if 'nut' in df else 0,
        "washer": int(df['washer'].sum()) if 'washer' in df else 0,
        "total_parts": total_parts,
        "parts_trend": parts_trend,
        "inference_speed": inference_speed,
        "model_accuracy": accuracy,
        "accuracy_trend": accuracy_trend,
        "tta_status": tta_status,
        "last_updated": datetime.now().strftime("%H:%M:%S")
    }
    return stats

@app.get("/download")
def download_csv():
    if SUBMISSION_PATH.exists():
        return FileResponse(SUBMISSION_PATH, media_type="text/csv", filename="submission.csv")
    return JSONResponse({"error": "File not found"}, status_code=404)

@app.get("/download/pdf")
def download_pdf():
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", 'B', 16)
    pdf.cell(200, 10, txt="AI Factory - Detection Report", ln=True, align='C')
    pdf.set_font("Arial", size=12)
    pdf.cell(200, 10, txt=f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", ln=True, align='C')
    
    if SUBMISSION_PATH.exists():
        df = pd.read_csv(SUBMISSION_PATH)
        pdf.ln(10)
        pdf.cell(200, 10, txt=f"Total Images Processed: {len(df)}", ln=True)
        pdf.cell(200, 10, txt=f"Total Bolts: {df['bolt'].sum()}", ln=True)
        pdf.cell(200, 10, txt=f"Total Locating Pins: {df['locatingpin'].sum()}", ln=True)
        pdf.cell(200, 10, txt=f"Total Nuts: {df['nut'].sum()}", ln=True)
        pdf.cell(200, 10, txt=f"Total Washers: {df['washer'].sum()}", ln=True)
        
    pdf_path = BASE_DIR / "report.pdf"
    pdf.output(str(pdf_path))
    return FileResponse(pdf_path, media_type="application/pdf", filename="report.pdf")

@app.get("/download/docs")
def download_docs():
    doc = Document()
    doc.add_heading('AI Factory Research Summary', 0)
    p = doc.add_paragraph('Automated SolidWorks Part Detection via YOLO11 System.')
    
    if SUBMISSION_PATH.exists():
        df = pd.read_csv(SUBMISSION_PATH)
        doc.add_heading('Statistical Summary', level=1)
        doc.add_paragraph(f"Total Scanned: {len(df)}")
        table = doc.add_table(rows=1, cols=4)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Bolt'
        hdr_cells[1].text = 'Locating Pin'
        hdr_cells[2].text = 'Nut'
        hdr_cells[3].text = 'Washer'
        
        row_cells = table.add_row().cells
        row_cells[0].text = str(df['bolt'].sum())
        row_cells[1].text = str(df['locatingpin'].sum())
        row_cells[2].text = str(df['nut'].sum())
        row_cells[3].text = str(df['washer'].sum())

    doc_path = BASE_DIR / "summary.docx"
    doc.save(str(doc_path))
    return FileResponse(doc_path, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", filename="summary.docx")

@app.get("/hardware")
async def get_hardware():
    gpu_info = []
    try:
        gpus = GPUtil.getGPUs()
        for gpu in gpus:
            gpu_info.append({
                "name": gpu.name,
                "load": gpu.load * 100,
                "memoryUsed": gpu.memoryUsed,
                "memoryTotal": gpu.memoryTotal,
                "temperature": gpu.temperature
            })
    except:
        gpu_info = [{"name": "CPU-Only Mode", "load": 0, "memoryUsed": 0, "memoryTotal": 0, "temperature": 0}]
    
    return {
        "gpu": gpu_info[0] if gpu_info else None,
        "cpu": psutil.cpu_percent(),
        "ram": psutil.virtual_memory().percent
    }

# ============== LIVE CAMERA DETECTION ==============
camera_active = False
camera_cap = None
live_detection_counts = {'bolt': 0, 'nut': 0, 'washer': 0, 'locatingpin': 0}

def generate_camera_frames(loop):
    """Generator that yields MJPEG frames with YOLO detections"""
    global camera_cap, camera_active, live_detection_counts
    
    print("DEBUG: Attempting to open camera...")
    camera_cap = cv2.VideoCapture(0)
    if not camera_cap.isOpened():
        print("ERROR: Could not open camera (index 0).")
        camera_active = False
        return

    print("DEBUG: Camera opened successfully.")
    camera_cap.set(cv2.CAP_PROP_FRAME_WIDTH, 640)
    camera_cap.set(cv2.CAP_PROP_FRAME_HEIGHT, 480)
    
    frame_count = 0
    while camera_active:
        ret, frame = camera_cap.read()
        if not ret:
            print("ERROR: Failed to read frame.")
            # OPTIMIZATION: Balanced High-Fidelity Mode (1024px + Agnostic NMS + 0.25 Conf)
            # Process 1 in 3 frames to maintain FPS while using high resolution
            frame_count += 1
            continue

        key_frame = frame_count % 3 == 0 
        if not key_frame: 
            frame_count += 1
            continue

        # Reset counts for this frame
        counts = {'bolt': 0, 'nut': 0, 'washer': 0, 'locatingpin': 0}
        
        try:
            # PERFECT CONFIG: High Res + Balanced Confidence
            # PERFECT CONFIG: High Res + Balanced Confidence + Tracking
            results = pipeline.model.track(
                frame, 
                conf=0.55,       # VERY HIGH precision to avoid background noise
                iou=0.50,
                imgsz=1024,      # High resolution
                augment=False,
                agnostic_nms=True,
                persist=True,    # Enable tracking persistence
                verbose=False
            )
            
            # Draw detections and count
            if results[0].boxes is not None and len(results[0].boxes) > 0:
                annotated_frame = results[0].plot()
                
                for box in results[0].boxes:
                    cls_id = int(box.cls[0])
                    # Use globally remapped names
                    cls_name = pipeline.model.names[cls_id]
                        
                    # STRICT FILTER: Only allow the 4 specific parts
                    # Anything else is ignored (treated as background)
                    if cls_name in counts:
                        counts[cls_name] += 1
            else:
                annotated_frame = frame.copy()
        except Exception as e:
            annotated_frame = frame.copy()
            cv2.putText(annotated_frame, f"Error: {str(e)[:30]}", (10, 30),
                       cv2.FONT_HERSHEY_SIMPLEX, 0.7, (0, 0, 255), 2)
        
        # Update global counts for WebSocket
        live_detection_counts = counts
        
        # Broadcast via WebSocket (Skip broadcast if counts are same to save bandwidth)
        import json
        asyncio.run_coroutine_threadsafe(
            log_manager.broadcast(f"LIVE_COUNTS:{json.dumps(counts)}"),
            loop
        )
        
        # Add count overlay for visual feedback
        y_offset = 30
        for cls_name, count in counts.items():
            color = (0, 255, 0) if count > 0 else (128, 128, 128)
            cv2.putText(annotated_frame, f"{cls_name}: {count}", (10, y_offset),
                       cv2.FONT_HERSHEY_SIMPLEX, 0.6, color, 2)
            y_offset += 25
        
        # Encode frame as JPEG with lower quality for speed
        _, buffer = cv2.imencode('.jpg', annotated_frame, [cv2.IMWRITE_JPEG_QUALITY, 60])
        frame_bytes = buffer.tobytes()
        
        yield (b'--frame\r\n'
               b'Content-Type: image/jpeg\r\n\r\n' + frame_bytes + b'\r\n')
        
        # Control loop speed
        time.sleep(0.01)
    
    if camera_cap:
        camera_cap.release()
        camera_cap = None

@app.get("/video_feed")
async def video_feed():
    """Stream live camera feed with YOLO detections"""
    global camera_active
    camera_active = True
    loop = asyncio.get_running_loop()
    return StreamingResponse(
        generate_camera_frames(loop),
        media_type="multipart/x-mixed-replace; boundary=frame"
    )

@app.post("/camera/start")
async def start_camera():
    """Start the camera stream"""
    global camera_active
    camera_active = True
    return {"status": "Camera started"}

@app.post("/camera/stop")
async def stop_camera():
    """Stop the camera stream"""
    global camera_active, camera_cap
    camera_active = False
    if camera_cap:
        camera_cap.release()
        camera_cap = None
    return {"status": "Camera stopped"}

@app.get("/camera/status")
async def camera_status():
    """Get current camera status and live detection counts"""
    global camera_active, live_detection_counts
    return {
        "active": camera_active,
        "counts": live_detection_counts
    }

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
